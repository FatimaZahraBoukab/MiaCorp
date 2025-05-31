from fastapi import APIRouter, Depends, HTTPException, Response
from typing import Optional
from ..couchdb import CouchDB
from .auth import get_current_active_user
import io
import re
import json
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter, A4
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, PageBreak
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY, TA_LEFT, TA_RIGHT
from reportlab.lib import colors
from reportlab.lib.units import inch, cm
from reportlab.platypus.flowables import HRFlowable
import logging
from ..utils.google_docs_utils import get_google_doc_content, extract_doc_id_from_url
from datetime import datetime

# Configuration du logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Pour les documents Word, utilisez python-docx
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.shared import OxmlElement, qn
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    logger.warning("python-docx not installed. Word export will not be available.")

router = APIRouter(
    prefix="/documents",
    tags=["documents"],
    responses={404: {"description": "Not found"}},
)

DOCUMENTS_COLLECTION = "documents"
ENTREPRISES_COLLECTION = "entreprises"
TEMPLATES_COLLECTION = "templates"

class DocumentLoopProcessor:
    """Processeur pour gérer les boucles dans les documents."""
    
    def __init__(self):
        self.loop_patterns = {
            'ACTIONNAIRES': r'{{#LOOP_ACTIONNAIRES}}(.*?){{/LOOP_ACTIONNAIRES}}',
            'GERANTS': r'{{#LOOP_GERANTS}}(.*?){{/LOOP_GERANTS}}',
            'PRESIDENTS': r'{{#LOOP_PRESIDENTS}}(.*?){{/LOOP_PRESIDENTS}}',
        }
        
        self.condition_patterns = {
            'IF_GERANT': r'{{#IF_GERANT}}(.*?){{/IF_GERANT}}',
            'IF_PRESIDENT': r'{{#IF_PRESIDENT}}(.*?){{/IF_PRESIDENT}}',
        }
    
    def process_document(self, content: str, variables: dict) -> str:
        """Traite un document complet avec toutes les boucles et conditions."""
        logger.info("Début du traitement des boucles dans le document")
        
        # Traiter les boucles
        for loop_type, pattern in self.loop_patterns.items():
            content = self._process_loop(content, pattern, loop_type, variables)
        
        # Traiter les conditions
        content = self._process_conditions(content, variables)
        
        logger.info("Traitement des boucles terminé")
        return content
    
    def _process_loop(self, content: str, pattern: str, loop_type: str, variables: dict) -> str:
        """Traite une boucle spécifique dans le document."""
        def replace_loop(match):
            loop_content = match.group(1).strip()
            logger.info(f"Traitement de la boucle {loop_type}")
            
            # Récupérer les données pour cette boucle
            loop_data = self._get_loop_data(loop_type, variables)
            
            if not loop_data:
                logger.warning(f"Aucune donnée trouvée pour la boucle {loop_type}")
                return ""
            
            logger.info(f"Trouvé {len(loop_data)} éléments pour la boucle {loop_type}")
            
            # Générer le contenu pour chaque élément
            result_parts = []
            for i, item_data in enumerate(loop_data, 1):
                item_content = loop_content
                
                # Remplacer les variables pour cet élément
                item_content = self._replace_item_variables(item_content, item_data, i)
                result_parts.append(item_content)
            
            return '\n\n'.join(result_parts)
        
        return re.sub(pattern, replace_loop, content, flags=re.DOTALL)
    
    def _get_loop_data(self, loop_type: str, variables: dict) -> list:
        """Récupère les données pour un type de boucle spécifique."""
        if loop_type == 'ACTIONNAIRES':
            return self._get_actionnaires_data(variables)
        elif loop_type == 'GERANTS':
            return self._get_gerants_data(variables)
        elif loop_type == 'PRESIDENTS':
            return self._get_presidents_data(variables)
        
        return []
    
    def _get_actionnaires_data(self, variables: dict) -> list:
        """Récupère les données des actionnaires."""
        # Essayer d'abord avec la liste JSON
        if 'liste_actionnaires' in variables:
            try:
                data = json.loads(variables['liste_actionnaires'])
                logger.info(f"Données actionnaires récupérées depuis liste_actionnaires: {len(data)} éléments")
                return data
            except (json.JSONDecodeError, TypeError) as e:
                logger.warning(f"Erreur lors du parsing de liste_actionnaires: {e}")
        
        # Sinon, extraire des variables individuelles
        return self._extract_indexed_data(variables, 'associe')
    
    def _get_gerants_data(self, variables: dict) -> list:
        """Récupère les données des gérants."""
        actionnaires = self._get_actionnaires_data(variables)
        # Pour l'instant, retourner tous les actionnaires comme gérants potentiels
        return actionnaires
    
    def _get_presidents_data(self, variables: dict) -> list:
        """Récupère les données des présidents (pour SAS)."""
        actionnaires = self._get_actionnaires_data(variables)
        # Pour l'instant, retourner tous les actionnaires comme présidents potentiels
        return actionnaires
    
    def _extract_indexed_data(self, variables: dict, suffix: str) -> list:
        """Extrait les données indexées des variables."""
        data_list = []
        
        # Trouver le nombre maximum d'index
        max_index = 0
        pattern = f'_{suffix}_(\d+)$'
        
        for key in variables.keys():
            match = re.search(pattern, key)
            if match:
                max_index = max(max_index, int(match.group(1)))
        
        logger.info(f"Index maximum trouvé pour {suffix}: {max_index}")
        
        # Si pas d'index trouvé, chercher les variables sans index
        if max_index == 0:
            base_fields = [
                f'nom_{suffix}', f'date_naissance_{suffix}', f'lieu_naissance_{suffix}',
                f'adresse_{suffix}', f'nationalite_{suffix}', f'apport_numeraire_{suffix}',
                f'nombre_parts_{suffix}', f'nombre_actions_{suffix}', f'nombre_actions_associe_{suffix}'
            ]
            
            item = {}
            for field in base_fields:
                if field in variables:
                    item[field] = variables[field]
            
            if item:
                data_list.append(item)
                logger.info("Données extraites sans index")
        else:
            # Extraire les données avec index
            for i in range(1, max_index + 1):
                item = {}
                base_fields = [
                    f'nom_{suffix}', f'date_naissance_{suffix}', f'lieu_naissance_{suffix}',
                    f'adresse_{suffix}', f'nationalite_{suffix}', f'apport_numeraire_{suffix}',
                    f'nombre_parts_{suffix}', f'nombre_actions_{suffix}', f'nombre_actions_associe_{suffix}'
                ]
                
                for field in base_fields:
                    key_with_index = f"{field}_{i}"
                    if key_with_index in variables:
                        item[field] = variables[key_with_index]
                
                if item:
                    data_list.append(item)
                    logger.info(f"Données extraites pour l'index {i}")
        
        return data_list
    
    def _replace_item_variables(self, content: str, item_data: dict, index: int) -> str:
        """Remplace les variables pour un élément spécifique."""
        # Remplacer les variables sans index
        for key, value in item_data.items():
            if key != 'id':
                placeholder = f"{{{{{key}}}}}"
                content = content.replace(placeholder, str(value))
        
        # Remplacer les variables avec index
        for key, value in item_data.items():
            if key != 'id':
                placeholder_with_index = f"{{{{{key}_{index}}}}}"
                content = content.replace(placeholder_with_index, str(value))
        
        return content
    
    def _process_conditions(self, content: str, variables: dict) -> str:
        """Traite les conditions dans le document."""
        for condition_type, pattern in self.condition_patterns.items():
            content = self._process_condition(content, pattern, condition_type, variables)
        
        return content
    
    def _process_condition(self, content: str, pattern: str, condition_type: str, variables: dict) -> str:
        """Traite une condition spécifique."""
        def replace_condition(match):
            condition_content = match.group(1)
            
            # Évaluer la condition (pour l'instant, toujours vraie)
            if True:  # Logique de condition à implémenter selon vos besoins
                return condition_content
            else:
                return ""
        
        return re.sub(pattern, replace_condition, content, flags=re.DOTALL)

@router.get("/export/{entreprise_id}")
async def export_document(
    entreprise_id: str,
    document_index: int = 0,
    format: str = "pdf",
    current_user: dict = Depends(get_current_active_user)
):
    """
    Exporte un document au format PDF ou DOCX avec les variables remplacées par les valeurs du client.
    Supporte maintenant les boucles d'actionnaires pour SARL et SAS.
    """
    # Validation et récupération des données
    entreprises_db = CouchDB(ENTREPRISES_COLLECTION)
    entreprise = await entreprises_db.get_by_id(entreprise_id)
    
    if not entreprise:
        raise HTTPException(status_code=404, detail="Entreprise non trouvée")
    
    logger.info(f"Entreprise trouvée: {entreprise_id}")
    logger.info(f"Template ID associé: {entreprise.get('template_id', 'Non défini')}")
    logger.info(f"Document index demandé: {document_index}")
    
    if current_user["role"] != "expert" and current_user["id"] != entreprise["client_id"]:
        raise HTTPException(status_code=403, detail="Vous n'êtes pas autorisé à accéder à ce document")
    
    if entreprise["statut"] != "validé":
        raise HTTPException(status_code=400, detail="L'entreprise doit être validée pour exporter le document")
    
    # Récupération du template
    if "template_id" not in entreprise or not entreprise["template_id"]:
        logger.warning(f"Pas de template_id dans l'entreprise {entreprise_id}, recherche par type")
        templates_db = CouchDB(TEMPLATES_COLLECTION)
        templates = await templates_db.query({"type_entreprise": entreprise["type"], "statut": "validé"})
        
        if not templates:
            raise HTTPException(
                status_code=404, 
                detail=f"Aucun template trouvé pour le type d'entreprise {entreprise['type']}"
            )
        
        template = templates[0]
        logger.info(f"Template trouvé par type: {template['id']}")
    else:
        templates_db = CouchDB(TEMPLATES_COLLECTION)
        template = await templates_db.get_by_id(entreprise["template_id"])
        
        if not template:
            logger.warning(f"Template {entreprise['template_id']} non trouvé, recherche par type")
            templates = await templates_db.query({"type_entreprise": entreprise["type"], "statut": "validé"})
            
            if not templates:
                raise HTTPException(
                    status_code=404, 
                    detail=f"Template {entreprise['template_id']} non trouvé et aucun template alternatif disponible"
                )
            
            template = templates[0]
            logger.info(f"Template alternatif trouvé: {template['id']}")
    
    if not template.get("documents") or len(template["documents"]) <= document_index:
        raise HTTPException(
            status_code=404, 
            detail=f"Document à l'index {document_index} non trouvé dans le template"
        )
    
    document = template["documents"][document_index]
    logger.info(f"Document sélectionné: {document.get('titre', 'Sans titre')}")
    
    # Récupérer le contenu du document depuis Google Docs
    try:
        doc_id = extract_doc_id_from_url(document.get('google_doc_id', ''))
        logger.info(f"Récupération du contenu du document depuis Google Docs: {doc_id}")
        doc_content = get_google_doc_content(doc_id)
        logger.info(f"Contenu récupéré avec succès, longueur: {len(doc_content)} caractères")
    except Exception as e:
        logger.error(f"Erreur lors de la récupération du contenu du document: {e}")
        raise HTTPException(status_code=500, detail=f"Erreur lors de la récupération du contenu du document: {str(e)}")
    
    # Traiter les boucles d'actionnaires si c'est un type SARL ou SAS
    valeurs_variables = entreprise["valeurs_variables"]
    
    if entreprise["type"] in ['SARL', 'SAS']:
        logger.info("Traitement des boucles d'actionnaires pour le type d'entreprise: " + entreprise["type"])
        processor = DocumentLoopProcessor()
        doc_content = processor.process_document(doc_content, valeurs_variables)
    
    # Remplacer les variables restantes par les valeurs
    doc_content_with_values = replace_variables(doc_content, valeurs_variables)
    
    # Générer le document au format demandé
    if format.lower() == "pdf":
        content, content_type = generate_beautiful_pdf(
            doc_content_with_values, 
            entreprise["type"], 
            document.get('titre', 'Document'), 
            entreprise
        )
        filename = f"{document.get('titre', 'document')}_{entreprise_id}.pdf"
    elif format.lower() == "docx":
        if not DOCX_AVAILABLE:
            raise HTTPException(status_code=501, detail="Export Word non disponible. Veuillez installer python-docx.")
        content, content_type = generate_beautiful_docx(
            doc_content_with_values, 
            entreprise["type"], 
            document.get('titre', 'Document'), 
            entreprise
        )
        filename = f"{document.get('titre', 'document')}_{entreprise_id}.docx"
    else:
        raise HTTPException(status_code=400, detail="Format non supporté. Utilisez 'pdf' ou 'docx'")
    
    return Response(
        content=content,
        media_type=content_type,
        headers={
            "Content-Disposition": f"attachment; filename={filename}"
        }
    )

def replace_variables(content: str, variables: dict) -> str:
    """
    Remplace les variables dans le contenu par leurs valeurs.
    Gère maintenant aussi les variables d'actionnaires avec index.
    """
    for key, value in variables.items():
        placeholder = "{{" + key + "}}"
        content = content.replace(placeholder, str(value))
    
    # Nettoyer les variables non remplacées
    remaining_vars = re.findall(r'{{(.*?)}}', content)
    
    for var in remaining_vars:
        # Ignorer les marqueurs de boucle et de condition
        if not var.startswith('#') and not var.startswith('/'):
            content = content.replace("{{" + var + "}}", "___________")
    
    return content

def create_custom_styles():
    """Crée des styles personnalisés pour les documents PDF."""
    styles = getSampleStyleSheet()
    
    # Style pour le titre principal
    title_style = ParagraphStyle(
        'CustomTitle',
        parent=styles['Heading1'],
        fontSize=24,
        spaceAfter=30,
        spaceBefore=20,
        textColor=colors.HexColor('#2C3E50'),
        alignment=TA_CENTER,
        fontName='Helvetica-Bold'
    )
    
    # Style pour les sous-titres
    subtitle_style = ParagraphStyle(
        'CustomSubtitle',
        parent=styles['Heading2'],
        fontSize=18,
        spaceAfter=20,
        spaceBefore=15,
        textColor=colors.HexColor('#34495E'),
        alignment=TA_LEFT,
        fontName='Helvetica-Bold'
    )
    
    # Style pour les sections
    section_style = ParagraphStyle(
        'CustomSection',
        parent=styles['Heading3'],
        fontSize=14,
        spaceAfter=12,
        spaceBefore=12,
        textColor=colors.HexColor('#2980B9'),
        alignment=TA_LEFT,
        fontName='Helvetica-Bold'
    )
    
    # Style pour le texte normal avec meilleur espacement
    normal_style = ParagraphStyle(
        'CustomNormal',
        parent=styles['Normal'],
        fontSize=11,
        spaceAfter=6,
        spaceBefore=6,
        textColor=colors.HexColor('#2C3E50'),
        alignment=TA_JUSTIFY,
        fontName='Helvetica',
        leading=14
    )
    
    return {
        'title': title_style,
        'subtitle': subtitle_style,
        'section': section_style,
        'normal': normal_style
    }

def parse_content_structure(content: str):
    """Parse le contenu pour identifier les titres, sections et paragraphes."""
    lines = content.split('\n')
    structured_content = []
    
    for line in lines:
        line = line.strip()
        if not line:
            continue
            
        # Détecter les titres (lignes en majuscules ou commençant par des mots-clés)
        if (line.isupper() and len(line) > 3) or any(line.upper().startswith(keyword) for keyword in ['TITRE', 'CHAPITRE', 'SECTION']):
            structured_content.append({'type': 'title', 'content': line})
        # Détecter les sous-titres (lignes avec première lettre en majuscule et certains patterns)
        elif (line[0].isupper() and ':' in line) or any(line.startswith(keyword) for keyword in ['Article', 'Clause', 'Paragraphe']):
            structured_content.append({'type': 'subtitle', 'content': line})
        # Détecter les sections (lignes commençant par des numéros)
        elif re.match(r'^\d+\.', line) or re.match(r'^[A-Z]\)', line) or line.startswith('ARTICLE'):
            structured_content.append({'type': 'section', 'content': line})
        else:
            structured_content.append({'type': 'normal', 'content': line})
    
    return structured_content

def generate_beautiful_pdf(content: str, entreprise_type: str, document_title: str, entreprise: dict) -> tuple:
    """Génère un document PDF magnifiquement formaté avec support des actionnaires multiples."""
    buffer = io.BytesIO()
    doc = SimpleDocTemplate(
        buffer, 
        pagesize=A4,
        rightMargin=2*cm,
        leftMargin=2*cm,
        topMargin=2.5*cm,
        bottomMargin=2*cm
    )
    
    # Créer les styles personnalisés
    custom_styles = create_custom_styles()
    elements = []
    
    # En-tête du document avec informations de l'entreprise
    header_data = [
       
        ['Type:', entreprise_type],
        ['Date:', datetime.now().strftime('%d/%m/%Y')],
        ['Document:', document_title]
    ]
    
    header_table = Table(header_data, colWidths=[3*cm, 6*cm])
    header_table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (0, -1), colors.HexColor('#ECF0F1')),
        ('TEXTCOLOR', (0, 0), (-1, -1), colors.HexColor('#2C3E50')),
        ('ALIGN', (0, 0), (-1, -1), 'LEFT'),
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTNAME', (1, 0), (1, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('GRID', (0, 0), (-1, -1), 1, colors.HexColor('#BDC3C7')),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
    ]))
    
    elements.append(header_table)
    elements.append(Spacer(1, 20))
    
    # Ligne de séparation décorative
    elements.append(HRFlowable(width="100%", thickness=2, color=colors.HexColor('#3498DB')))
    elements.append(Spacer(1, 20))
    
    # Titre principal du document
    elements.append(Paragraph(document_title, custom_styles['title']))
    elements.append(Spacer(1, 10))
    
    # Ligne de séparation sous le titre
    elements.append(HRFlowable(width="60%", thickness=1, color=colors.HexColor('#95A5A6')))
    elements.append(Spacer(1, 20))
    
    # Parser et formater le contenu
    structured_content = parse_content_structure(content)
    
    for item in structured_content:
        content_type = item['type']
        content_text = item['content']
        
        try:
            if content_type == 'title':
                elements.append(Spacer(1, 15))
                elements.append(Paragraph(content_text, custom_styles['title']))
                elements.append(HRFlowable(width="100%", thickness=1, color=colors.HexColor('#3498DB')))
                elements.append(Spacer(1, 10))
                
            elif content_type == 'subtitle':
                elements.append(Spacer(1, 12))
                elements.append(Paragraph(content_text, custom_styles['subtitle']))
                elements.append(Spacer(1, 8))
                
            elif content_type == 'section':
                elements.append(Paragraph(content_text, custom_styles['section']))
                elements.append(Spacer(1, 6))
                
            else:  # normal
                elements.append(Paragraph(content_text, custom_styles['normal']))
                
        except Exception as e:
            logger.warning(f"Erreur lors du formatage d'un paragraphe: {e}")
            elements.append(Paragraph(f"[Contenu: {content_text[:50]}...]", custom_styles['normal']))
    
    # Pied de page avec ligne de séparation
    elements.append(Spacer(1, 30))
    elements.append(HRFlowable(width="100%", thickness=2, color=colors.HexColor('#3498DB')))
    elements.append(Spacer(1, 10))
    
    footer_style = ParagraphStyle(
        'Footer',
        fontSize=9,
        textColor=colors.HexColor('#7F8C8D'),
        alignment=TA_CENTER
    )
    elements.append(Paragraph("Document généré automatiquement", footer_style))
    
    # Construire le document
    doc.build(elements)
    
    pdf_content = buffer.getvalue()
    buffer.close()
    
    return pdf_content, "application/pdf"

def generate_beautiful_docx(content: str, entreprise_type: str, document_title: str, entreprise: dict) -> tuple:
    """Génère un document DOCX magnifiquement formaté avec support des actionnaires multiples."""
    doc = Document()
    
    # Titre principal
    title = doc.add_heading(document_title, 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Informations de l'entreprise
    info_para = doc.add_paragraph()
    info_para.add_run(f"Type d'entreprise: ").bold = True
    info_para.add_run(f"{entreprise_type}\n")
    info_para.add_run(f"Nom: ").bold = True
    info_para.add_run(f"{entreprise.get('nom', 'N/A')}\n")
    info_para.add_run(f"Date de génération: ").bold = True
    info_para.add_run(f"{datetime.now().strftime('%d/%m/%Y')}")
    
    doc.add_paragraph()  # Espacement
    
    # Contenu principal
    paragraphs = content.split('\n')
    for paragraph in paragraphs:
        if paragraph.strip():
            doc.add_paragraph(paragraph.strip())
    
    # Enregistrer dans un buffer
    buffer = io.BytesIO()
    doc.save(buffer)
    
    docx_content = buffer.getvalue()
    buffer.close()
    
    return docx_content, "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

# Endpoints de debug et utilitaires
@router.get("/debug/{entreprise_id}")
async def debug_document_export(
    entreprise_id: str,
    current_user: dict = Depends(get_current_active_user)
):
    """Endpoint de diagnostic pour aider à déboguer les problèmes d'exportation de documents."""
    if current_user["role"] != "expert" and current_user["role"] != "admin":
        raise HTTPException(status_code=403, detail="Seuls les experts et administrateurs peuvent accéder à cette fonctionnalité")
    
    entreprises_db = CouchDB(ENTREPRISES_COLLECTION)
    entreprise = await entreprises_db.get_by_id(entreprise_id)
    
    if not entreprise:
        return {"error": "Entreprise non trouvée", "entreprise_id": entreprise_id}
    
    template_id = entreprise.get("template_id")
    template_info = {"template_id": template_id}
    
    if not template_id:
        template_info["error"] = "Pas de template_id dans l'entreprise"
    else:
        templates_db = CouchDB(TEMPLATES_COLLECTION)
        template = await templates_db.get_by_id(template_id)
        
        if not template:
            template_info["error"] = f"Template {template_id} non trouvé dans la base de données"
        else:
            template_info["found"] = True
            template_info["titre"] = template.get("titre")
            template_info["type_entreprise"] = template.get("type_entreprise")
            template_info["statut"] = template.get("statut")
            template_info["supports_dynamic_shareholders"] = template.get("supports_dynamic_shareholders", False)
            
            if template.get("documents"):
                template_info["documents"] = []
                for i, doc in enumerate(template["documents"]):
                    doc_info = {
                        "index": i,
                        "titre": doc.get("titre"),
                        "google_doc_id": doc.get("google_doc_id"),
                        "has_shareholder_loop": doc.get("has_shareholder_loop", False)
                    }
                    template_info["documents"].append(doc_info)
    
    templates_db = CouchDB(TEMPLATES_COLLECTION)
    available_templates = await templates_db.query({"type_entreprise": entreprise["type"], "statut": "validé"})
    
    return {
        "entreprise": {
            "id": entreprise["id"],
            "type": entreprise["type"],
            "statut": entreprise["statut"],
            "template_id": entreprise.get("template_id")
        },
        "template_info": template_info,
        "available_templates": [
            {"id": t["id"], "titre": t["titre"], "supports_dynamic_shareholders": t.get("supports_dynamic_shareholders", False)} 
            for t in available_templates
        ],
        "valeurs_variables": entreprise.get("valeurs_variables", {}),
        "actionnaires_data": entreprise.get("valeurs_variables", {}).get("liste_actionnaires", "Non trouvé")
    }

@router.get("/test-google-doc/{doc_id}")
async def test_google_doc_access(
    doc_id: str,
    current_user: dict = Depends(get_current_active_user)
):
    """Teste l'accès à un document Google Docs et retourne un aperçu du contenu."""
    if current_user["role"] != "expert" and current_user["role"] != "admin":
        raise HTTPException(status_code=403, detail="Seuls les experts et administrateurs peuvent accéder à cette fonctionnalité")
    
    try:
        doc_id = extract_doc_id_from_url(doc_id)
        content = get_google_doc_content(doc_id)
        
        # Détecter les boucles dans le contenu
        has_loops = bool(re.search(r'{{#LOOP_.*?}}', content))
        loop_types = re.findall(r'{{#LOOP_(.*?)}}', content)
        
        return {
            "status": "success",
            "doc_id": doc_id,
            "content_length": len(content),
            "content_preview": content[:500] + "..." if len(content) > 500 else content,
            "has_loops": has_loops,
            "loop_types": loop_types
        }
    except Exception as e:
        logger.error(f"Erreur lors du test d'accès au document Google Docs: {e}")
        return {
            "status": "error",
            "doc_id": doc_id,
            "error": str(e)
        }

@router.get("/available/{entreprise_id}")
async def get_available_documents(
    entreprise_id: str,
    current_user: dict = Depends(get_current_active_user)
):
    """Retourne la liste des documents disponibles pour une entreprise."""
    entreprises_db = CouchDB(ENTREPRISES_COLLECTION)
    entreprise = await entreprises_db.get_by_id(entreprise_id)
    
    if not entreprise:
        raise HTTPException(status_code=404, detail="Entreprise non trouvée")
    
    if current_user["role"] != "expert" and current_user["id"] != entreprise["client_id"]:
        raise HTTPException(status_code=403, detail="Vous n'êtes pas autorisé à accéder à ce document")
    
    if entreprise["statut"] != "validé":
        raise HTTPException(status_code=400, detail="L'entreprise doit être validée pour accéder aux documents")
    
    template_id = entreprise.get("template_id")
    
    if not template_id:
        templates_db = CouchDB(TEMPLATES_COLLECTION)
        templates = await templates_db.query({"type_entreprise": entreprise["type"], "statut": "validé"})
        
        if not templates:
            return {"documents": []}
        
        template = templates[0]
    else:
        templates_db = CouchDB(TEMPLATES_COLLECTION)
        template = await templates_db.get_by_id(template_id)
        
        if not template:
            templates = await templates_db.query({"type_entreprise": entreprise["type"], "statut": "validé"})
            
            if not templates:
                return {"documents": []}
            
            template = templates[0]
    
    return {
        "documents": [
            {
                "index": idx,
                "titre": doc.get("titre", f"Document {idx+1}"),
                "description": doc.get("description", ""),
                "has_shareholder_loop": doc.get("has_shareholder_loop", False)
            }
            for idx, doc in enumerate(template.get("documents", []))
        ],
        "supports_dynamic_shareholders": template.get("supports_dynamic_shareholders", False)
    }
